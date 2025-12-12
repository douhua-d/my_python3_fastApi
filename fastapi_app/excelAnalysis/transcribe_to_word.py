import whisper
from docx import Document
import os
import warnings
import datetime

# Filter warnings to keep output clean
warnings.filterwarnings("ignore")

def format_timestamp(seconds):
    """Convert seconds to MM:SS format"""
    seconds = int(seconds)
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    else:
        return f"{minutes:02d}:{secs:02d}"

def transcribe_and_save(audio_path, output_docx_path):
    print(f"Loading Whisper model (this may take a moment)...")
    # Using 'small' model. 
    model = whisper.load_model("small")
    
    print(f"Transcribing {audio_path}...")
    # Transcribe the audio
    result = model.transcribe(audio_path)
    
    print("Transcription complete. Saving to Word...")
    doc = Document()
    doc.add_heading('语音转写记录', 0)
    
    # Iterate through segments to provide better formatting with timestamps
    for segment in result["segments"]:
        start_time = format_timestamp(segment["start"])
        end_time = format_timestamp(segment["end"])
        text = segment["text"].strip()
        
        # Add timestamp and text as a paragraph
        # Create a paragraph
        p = doc.add_paragraph()
        
        # Add bold timestamp
        timestamp_run = p.add_run(f"[{start_time} - {end_time}] ")
        timestamp_run.bold = True
        
        # Add text
        text_run = p.add_run(text)
        
    doc.save(output_docx_path)
    print(f"Saved to {output_docx_path}")

if __name__ == "__main__":
    # The specific file requested
    audio_file = "15712999190(15712999190)_20251212165312.mp3"
    
    # Determine output filename
    docx_file = os.path.splitext(audio_file)[0] + ".docx"
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    audio_path = os.path.join(current_dir, audio_file)
    output_path = os.path.join(current_dir, docx_file)
    
    if os.path.exists(audio_path):
        try:
            transcribe_and_save(audio_path, output_path)
        except Exception as e:
            print(f"An error occurred: {e}")
    else:
        print(f"File {audio_file} not found in {current_dir}")
