from docx import Document
from docx.shared import Pt
import os

def add_summary_to_docx(docx_path, summary_text):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    
    # Create a new document to hold the content with summary at the top
    new_doc = Document()
    
    # Add Title
    new_doc.add_heading('语音转写记录', 0)
    
    # Add Summary Section
    new_doc.add_heading('主要内容总结', level=1)
    summary_para = new_doc.add_paragraph(summary_text)
    
    # Add a separator or space
    new_doc.add_paragraph()
    new_doc.add_heading('转写详情', level=1)
    
    # Copy existing content (skipping the first heading if it was '语音转写记录')
    # We need to be careful not to just copy text but runs to preserve formatting (bold timestamps)
    
    first_heading_skipped = False
    for element in doc.element.body:
        if element.tag.endswith('p'): # Paragraph
            # We can't easily access the paragraph object from the element directly in a way that preserves exact formatting 
            # without re-constructing it. 
            # Simpler approach: Iterate paragraphs and reconstruct.
            pass
            
    # Actually, python-docx doesn't make "insert at beginning" easy. 
    # It's easier to create a new doc, add summary, then copy content.
    # But copying content preserving formatting is tricky.
    
    # Alternative: Insert paragraph at index 0?
    # doc.paragraphs.insert(0, ...) is not directly supported in the API cleanly.
    
    # Let's try to manipulate the element tree to insert before the first paragraph.
    # But first, let's just use the `doc` object.
    
    # We will insert the summary *after* the title (index 0 usually).
    # Assuming first paragraph is title? No, `add_heading` adds a paragraph with style.
    
    # Let's simple append the summary to a new document, then iterate the old document's paragraphs 
    # and copy them to the new document.
    
    for para in doc.paragraphs:
        # Skip the original title if we added a new one, or just keep it?
        # The original doc has '语音转写记录' as title.
        if para.text == '语音转写记录' and para.style.name.startswith('Heading'):
            continue
            
        new_p = new_doc.add_paragraph()
        new_p.style = para.style
        
        for run in para.runs:
            new_r = new_p.add_run(run.text)
            new_r.bold = run.bold
            new_r.italic = run.italic
            new_r.underline = run.underline
            new_r.font.name = run.font.name
            # Copy other font attributes if needed
            
    new_doc.save(docx_path)
    print(f"Summary added to {docx_path}")

if __name__ == "__main__":
    file_path = "fastapi_app/excelAnalysis/15712999190(15712999190)_20251212165312.docx"
    abs_path = os.path.abspath(file_path)
    
    summary = """本录音为医药代表与李鹏飞经理关于药品（喜纳卡塞、依福卡塞）进院事宜的沟通。

1. 涉及医院：房山第一医院、房山中医院。
2. 房山第一医院情况：
   - 肾内科主任对进口药较为敏感，建议找药剂科。
   - 药剂科主任名为刘洋（男）。
   - 进药需通过药事会，约一年召开一次，会前约半个月有通知。
   - 目前无该药销售。
3. 房山中医院情况：
   - 流程与房山第一医院类似，需走正规流程。
   - 建议尝试“临采”（临时采购）作为突破口。
   - 临床反馈较房山第一医院积极。
4. 后续：双方互加微信保持联系。"""

    # We need to do a better job of inserting. 
    # Instead of full copy which might lose some details, 
    # let's try to insert paragraphs at the beginning using internal API.
    
    doc = Document(abs_path)
    
    # Create paragraphs for summary
    # We want to insert them after the Title (index 0). 
    # Paragraphs are in doc.paragraphs.
    # The body elements are in doc.element.body.
    
    # Let's generate the summary elements
    heading = doc.add_heading('主要内容总结', level=1)
    para = doc.add_paragraph(summary)
    heading_2 = doc.add_heading('转写详情', level=1)
    
    # Now move them to the top (after the main title)
    # Assuming the first element is the main title.
    
    body = doc.element.body
    # The last 3 elements are what we just added.
    summary_elements = body[-3:]
    
    # Remove them from end
    for _ in range(3):
        body.remove(body[-1])
        
    # Insert them at index 1 (after title)
    # If there is no title, index 0.
    
    insert_index = 0
    if len(doc.paragraphs) > 0 and doc.paragraphs[0].style.name.startswith('Heading'):
        insert_index = 1
        
    # We need to find the element corresponding to paragraph at insert_index
    # This is tricky because paragraphs list only contains paragraphs, not tables etc.
    # But here we only have paragraphs.
    
    ref_element = None
    if len(body) > insert_index:
        ref_element = body[insert_index]
    
    for i, elem in enumerate(summary_elements):
        if ref_element:
            body.insert(body.index(ref_element), elem)
        else:
            body.append(elem)
            
    doc.save(abs_path)
    print(f"Summary inserted into {abs_path}")
